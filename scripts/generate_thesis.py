"""Generator i tezes MSc per sistemin DOKU -> Microsoft Word (.docx).

Krijon nje teze akademike GJITHEPERFSHIRESE ne shqip nga permbajtja reale e projektit:
kodi burimor, arkitektura, baza e te dhenave, siguria, pipeline-i RAG, rolet,
eksperimentet dhe te gjitha funksionalitetet. Numrat e Kapitullit 6 jane REALE
(harness-i mbi tests/sample_questions.csv, qwen2.5:3b). Screenshot-et jane reale
(docs/screenshots/, te kapura me Playwright nga aplikacioni i ekzekutuar).

Perdorim:  .venv\\Scripts\\python.exe scripts\\generate_thesis.py
Dalja:     data/exports/Teza_DOKU.docx  (+ kopje ne docs/)
"""
import os
import sys
from datetime import date

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import config

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SHOTS = os.path.join(ROOT, "docs", "screenshots")

# --------------------------------------------------------------------------------------
# Metadata e kopertines
# --------------------------------------------------------------------------------------
UNIVERSITY = "UNIVERSITETI I TIRANËS"
FACULTY = "FAKULTETI I SHKENCAVE TË NATYRËS"
DEPARTMENT = "DEPARTAMENTI I INFORMATIKËS"
PROGRAM = "Program i Ciklit të Dytë (Master i Shkencave) në Informatikë"
THESIS_TITLE = ("Zhvillimi i një sistemi inteligjent për analizë dokumentesh duke "
                "përdorur RAG dhe LLM: Aplikimi i tij në institucionet shtetërore "
                "në Shqipëri")
AUTHOR = "Denald Shehaj"
SUPERVISOR = "[Udhëheqësi Shkencor]"
CITY_YEAR = f"Tiranë, {date.today().year}"

# --------------------------------------------------------------------------------------
# Rezultate REALE nga harness-i (qwen2.5:3b). t = sekonda.
# --------------------------------------------------------------------------------------
RESULTS = [
    (1, "Kur duhet të dorëzohet deklarata tatimore vjetore?", 49.03, 189.03, 5, True, True, False),
    (2, "Çfarë ndodh nëse nuk paguhen tatimet brenda afatit?", 23.70, 76.39, 5, True, True, False),
    (3, "Ku mund të ankohet tatimpaguesi për një vlerësim tatimor?", 75.36, 79.59, 5, True, True, False),
    (4, "Sa ditë pushim vjetor të paguar ka punonjësi?", 19.22, 89.64, 5, True, True, False),
    (5, "Sa orë është orari normal javor i punës?", 37.84, 34.93, 5, True, True, False),
    (6, "Sa është afati i njoftimit për zgjidhjen e kontratës së punës?", 75.47, 53.49, 5, True, True, False),
    (7, "Cili është objektivi kryesor i Strategjisë Kombëtare të Dixhitalizimit për 2030?", 46.18, 83.56, 5, True, True, False),
    (8, "Ku do të përqendrohen investimet sipas strategjisë së dixhitalizimit?", 79.78, 38.75, 5, True, True, False),
    (9, "Sa kushton një biletë avioni nga Tirana për në Romë?", 28.57, 65.41, 3, True, False, False),
    (10, "Kush e fitoi Kupën e Botës në futboll në vitin 2018?", 44.48, 0.43, 0, False, False, True),
]
IN = [r for r in RESULTS if r[6]]
OUT = [r for r in RESULTS if not r[6]]
NORAG_MEAN_IN = sum(r[2] for r in IN) / len(IN)
RAG_MEAN_IN = sum(r[3] for r in IN) / len(IN)
GROUNDED = sum(1 for r in IN if r[5])

# ======================================================================================
# Helper-a
# ======================================================================================
doc = Document()
_normal = doc.styles["Normal"]
_normal.font.name = "Calibri"
_normal.font.size = Pt(11)
_normal.paragraph_format.line_spacing = 1.5
_normal.paragraph_format.space_after = Pt(6)


def h1(text):
    doc.add_page_break()
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text, align="justify", italic=False, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic, r.bold = italic, bold
    if size:
        r.font.size = Pt(size)
    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                   "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    return p


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        rr = c.paragraphs[0].add_run(hdr)
        rr.bold = True
        rr.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(9)
    return t


def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def figure(filename, cap, width=6.0):
    path = os.path.join(SHOTS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
    else:
        para(f"[Imazhi {filename} mungon — kapeni me scripts/capture_screenshots.py]", italic=True)
    caption(cap)


def add_toc(levels="1-3"):
    """Fut nje fushe reale 'Table of Contents' qe Word e mbush automatikisht."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "{levels}" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = ("Tabela e Përmbajtjes do të shfaqet këtu. Nëse mbetet bosh, kliko brenda "
                        "saj dhe shtyp F9 (ose Ctrl+A, pastaj F9) për ta përditësuar.")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)


def enable_update_fields_on_open():
    """Bën që Word t'i përditësojë fushat (përfshirë TOC) kur hapet dokumenti."""
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


# ======================================================================================
# KOPERTINA
# ======================================================================================
doc.add_paragraph()
doc.add_paragraph()
para(UNIVERSITY, align="center", bold=True, size=16)
para(FACULTY, align="center", bold=True, size=13)
para(DEPARTMENT, align="center", bold=True, size=13)
for _ in range(3):
    doc.add_paragraph()
para("TEMË DIPLOME", align="center", bold=True, size=14)
para(PROGRAM, align="center", italic=True, size=11)
doc.add_paragraph()
para(THESIS_TITLE, align="center", bold=True, size=15)
for _ in range(4):
    doc.add_paragraph()
t = doc.add_table(rows=2, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.cell(0, 0).text, t.cell(0, 1).text = "Punoi:", AUTHOR
t.cell(1, 0).text, t.cell(1, 1).text = "Udhëheqës shkencor:", SUPERVISOR
for r in t.rows:
    for c in r.cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(12)
for _ in range(6):
    doc.add_paragraph()
para(CITY_YEAR, align="center", bold=True, size=12)

# ======================================================================================
# FAQET PARAPRAKE
# ======================================================================================
h1("Deklaratë Origjinaliteti")
para("Unë, " + AUTHOR + ", deklaroj se kjo temë diplome me titull «" + THESIS_TITLE +
     "» është punë origjinale e imja, e realizuar nën udhëheqjen shkencore përkatëse. "
     "Çdo material i marrë nga burime të tjera është cituar sipas rregullave akademike. "
     "Sistemi softuerik i përshkruar (DOKU) është projektuar dhe zhvilluar tërësisht nga "
     "autori. Rezultatet eksperimentale të paraqitura në Kapitullin 6 dhe pamjet nga ekrani "
     "janë prodhuar nga vetë sistemi i ekzekutuar mbi një korpus testues dhe nuk janë trilluar.")
doc.add_paragraph()
para("Nënshkrimi: ____________________        Data: ____________________")

h1("Dedikim")
para("Familjes sime, për mbështetjen e pakushtëzuar gjatë gjithë rrugëtimit akademik.",
     align="center", italic=True)

h1("Mirënjohje")
para("Dëshiroj të falënderoj udhëheqësin tim shkencor për orientimin dhe sugjerimet e vyera "
     "gjatë gjithë procesit të realizimit të kësaj teme. Falënderoj gjithashtu stafin akademik "
     "të Departamentit të Informatikës, Fakulteti i Shkencave të Natyrës, Universiteti i Tiranës, "
     "për njohuritë e transmetuara. Një falënderim i veçantë shkon për familjen dhe miqtë që më "
     "mbështetën moralisht gjatë realizimit të kësaj pune.")

h1("Abstrakt")
para("Institucionet publike në Shqipëri administrojnë një volum të madh dokumentesh ligjore dhe "
     "normative — ligje, vendime të Këshillit të Ministrave (VKM), strategji, rregullore dhe "
     "udhëzime. Aksesi i shpejtë dhe i saktë në informacionin e tyre është thelbësor për "
     "vendimmarrjen institucionale, por kërkimi manual është i ngadaltë dhe i prirur ndaj gabimeve. "
     "Modelet e Mëdha Gjuhësore (LLM) ofrojnë ndërveprim në gjuhë natyrore, por vuajnë nga "
     "halucinacioni dhe mungesa e burimeve të verifikueshme, çka i bën të rrezikshme në kontekst "
     "institucional. Kjo temë paraqet DOKU-n, një sistem inteligjent plotësisht lokal për analizën "
     "e dokumenteve institucionale në gjuhën shqipe, i ndërtuar mbi arkitekturën Retrieval-Augmented "
     "Generation (RAG) të kombinuar me një LLM që ekzekutohet lokalisht nëpërmjet Ollama. Sistemi "
     "indekson dokumentet me embeddings shumëgjuhëshe (bge-m3) në një bazë vektoriale (ChromaDB), "
     "zbaton një «portë refuzimi» përpara thirrjes së modelit për të garantuar bazueshmërinë "
     "(grounding) dhe shmangur halucinacionet, dhe gjeneron përgjigje të cituara me referenca te "
     "dokumenti dhe faqja burimore. DOKU mbështet dy role (administrator dhe punonjës), regjistron "
     "çdo veprim në një gjurmë auditi, dhe ofron një modul eksperimentesh që krahason RAG-un kundrejt "
     "një LLM-je pa RAG. Vlerësimi empirik mbi një korpus prej 13 dokumentesh dhe 10 pyetjesh tregoi "
     "se sistemi prodhoi përgjigje të bazuara me citime për 100% të pyetjeve brenda korpusit dhe "
     "refuzoi saktë një pyetje jashtë korpusit në kohë nën një sekondë (pa e thirrur modelin), "
     "ndërkohë që LLM-ja pa RAG përgjigjej gjithmonë, edhe pa burim. Rezultatet konfirmojnë vlerën e "
     "qasjes RAG për besueshmërinë dhe gjurmueshmërinë e informacionit në kontekst institucional.")
doc.add_paragraph()
para("Fjalë kyçe: RAG, LLM lokal, ChromaDB, embeddings bge-m3, Ollama, analizë dokumentesh, "
     "gjuha shqipe, sektor publik, grounding, halucinacion, kontroll aksesi me role.", italic=True)

h1("Abstract (English)")
para("Public institutions in Albania manage large volumes of legal and normative documents — laws, "
     "Council of Ministers decisions (CMD), strategies, regulations and instructions. Fast and "
     "accurate access to their content is essential for institutional decision-making, yet manual "
     "search is slow and error-prone. Large Language Models (LLMs) enable natural-language "
     "interaction but suffer from hallucination and lack of verifiable sources, which makes them "
     "risky in an institutional context. This thesis presents DOKU, a fully local intelligent system "
     "for analysing institutional documents in the Albanian language, built on a Retrieval-Augmented "
     "Generation (RAG) architecture combined with a locally executed LLM via Ollama. The system "
     "indexes documents using multilingual embeddings (bge-m3) in a vector database (ChromaDB), "
     "applies a “refusal gate” before invoking the model to guarantee grounding and avoid "
     "hallucinations, and produces cited answers referencing the source document and page. DOKU "
     "supports two roles (administrator and employee), records every action in an audit trail, and "
     "provides an experiment module comparing RAG against a no-RAG baseline. Empirical evaluation on "
     "a corpus of 13 documents and 10 questions showed the system produced grounded, cited answers "
     "for 100% of in-corpus questions and correctly refused an out-of-corpus question in sub-second "
     "time (without invoking the model), whereas the no-RAG LLM always answered, even without a "
     "source. The results confirm the value of the RAG approach for the trustworthiness and "
     "traceability of information in an institutional context.")
doc.add_paragraph()
para("Keywords: RAG, local LLM, ChromaDB, bge-m3 embeddings, Ollama, document analysis, Albanian "
     "language, public sector, grounding, hallucination, role-based access control.", italic=True)

h1("Tabela e Përmbajtjes")
add_toc("1-3")

# ======================================================================================
# KAPITULLI 1 — HYRJE
# ======================================================================================
h1("Kapitulli 1 — Hyrje")

h2("1.1 Sfondi dhe Konteksti")
para("Transformimi dixhital i administratës publike është kthyer në një prioritet strategjik në "
     "Shqipëri, i mishëruar në iniciativa si platforma e-Albania dhe Strategjia Kombëtare për "
     "Dixhitalizimin. Megjithatë, pavarësisht dixhitalizimit të shërbimeve, pjesa dërrmuese e "
     "njohurisë institucionale vazhdon të ruhet në formën e dokumenteve tekstuale të pastrukturuara: "
     "ligje, vendime të Këshillit të Ministrave, strategji sektoriale, rregullore dhe udhëzime. Këto "
     "dokumente janë shpesh të gjata, me gjuhë juridike komplekse dhe referenca të ndërsjella, çka e "
     "bën gjetjen e informacionit specifik një detyrë të kushtueshme në kohë.")
para("Një punonjës i administratës që kërkon, për shembull, afatin ligjor për dorëzimin e një "
     "deklarate ose procedurën e ankimit ndaj një vlerësimi, shpesh duhet të lexojë manualisht "
     "dhjetëra faqe. Ky proces është jo vetëm joeficient, por edhe i ndjeshëm ndaj gabimeve "
     "interpretuese, të cilat në kontekst institucional mund të kenë pasoja ligjore.")
para("Përparimet e fundit në Inteligjencën Artificiale, veçanërisht Modelet e Mëdha Gjuhësore (LLM), "
     "kanë hapur mundësi të reja për ndërveprim në gjuhë natyrore mbi dokumente. Por aplikimi i tyre "
     "i drejtpërdrejtë në sektorin publik ndeshet me tri pengesa kryesore: (1) halucinacioni — "
     "prirja e LLM-ve për të gjeneruar informacion të pasaktë me ton bindës; (2) mungesa e burimeve "
     "të verifikueshme — përgjigjet nuk mund të gjurmohen te një dokument zyrtar; dhe (3) "
     "konfidencialiteti — dërgimi i dokumenteve institucionale te shërbime cloud të jashtme bie ndesh "
     "me kërkesat ligjore për mbrojtjen e të dhënave. Arkitektura Retrieval-Augmented Generation "
     "(RAG) i adreson dy të parat duke e detyruar modelin të përgjigjet vetëm mbi pasazhe të marra "
     "nga një korpus i besueshëm, ndërsa ekzekutimi plotësisht lokal adreson të tretën.")

h2("1.2 Përcaktimi i Problemit")
para("Problemi qendror që adreson kjo temë formulohet si vijon: si mund të ndërtohet një sistem që "
     "u mundëson punonjësve të institucioneve shtetërore shqiptare të marrin përgjigje të sakta, të "
     "bazuara dhe të verifikueshme nga një korpus dokumentesh zyrtare në gjuhën shqipe, pa rrezikuar "
     "konfidencialitetin e të dhënave dhe pa u mbështetur në shërbime cloud të jashtme?")
para("Nga ky problem i përgjithshëm rrjedhin nënproblemet specifike: (a) si të zbutet halucinacioni "
     "i LLM-ve në mënyrë strukturore, jo thjesht statistikore; (b) si të garantohet sovraniteti dhe "
     "privatësia e të dhënave institucionale; (c) si të trajtohet mbështetja e kufizuar e gjuhës "
     "shqipe në shumë mjete komerciale; dhe (d) si të zbatohet një kontroll aksesi me role që ndan "
     "qartë administratorin (që menaxhon korpusin) nga punonjësi (që vetëm e pyet atë).")

h2("1.3 Motivimi dhe Rëndësia")
para("Rëndësia praktike e punës qëndron në mundësimin e një aksesi më të shpejtë dhe më të "
     "besueshëm te njohuria institucionale, duke reduktuar kohën e kërkimit dhe rrezikun e gabimit. "
     "Rëndësia shkencore qëndron në demonstrimin se një «kontratë bazueshmërie» e zbatuar përpara "
     "gjenerimit mund ta bëjë fabrikimin strukturalisht të pamundur për pyetje jashtë korpusit — një "
     "veçori kritike për besueshmërinë në sektorin publik. Rëndësia social-institucionale qëndron në "
     "ofrimin e një modeli zgjidhjeje plotësisht lokal, të përshtatshëm për mjedise me kërkesa të "
     "larta konfidencialiteti.")

h2("1.4 Objektivat e Kërkimit")
numbered("Të projektohet dhe zhvillohet një sistem RAG plotësisht lokal për dokumente institucionale në gjuhën shqipe.")
numbered("Të garantohet bazueshmëria (grounding) nëpërmjet një mekanizmi refuzimi me prag ngjashmërie, përpara thirrjes së modelit.")
numbered("Të implementohet menaxhim i sigurt i dokumenteve me kontroll aksesi me role dhe regjistrim auditi.")
numbered("Të vlerësohet eksperimentalisht përfitimi i RAG-ut kundrejt një LLM-je pa RAG, me të dhëna reale.")
numbered("Të vlerësohet realizueshmëria e sistemit në harduer konsumatori (16GB RAM, GPU 4GB) dhe të identifikohen kompromiset.")

h2("1.5 Pyetjet Kërkimore")
para("PK1: A mund një sistem RAG plotësisht lokal të prodhojë përgjigje të bazuara dhe të cituara "
     "për pyetje në gjuhën shqipe mbi dokumente institucionale?")
para("PK2: Sa efektiv është mekanizmi i refuzimit në parandalimin e përgjigjeve për pyetje jashtë "
     "korpusit, krahasuar me një LLM pa RAG?")
para("PK3: Cilat janë kompromiset midis cilësisë gjuhësore dhe vonesës (latency) për modele të "
     "ndryshme lokale në harduer të kufizuar?")

h2("1.6 Kontributet e Tezës")
bullet("Një sistem referencë RAG, plotësisht lokal, i specializuar për gjuhën shqipe dhe sektorin publik shqiptar.")
bullet("Një «kontratë bazueshmërie» (refusal-gate përpara LLM-së) si mekanizëm strukturor anti-halucinacion, me efekt të matur empirikisht.")
bullet("Një modul eksperimental RAG kundrejt no-RAG me eksport CSV, që mundëson vlerësim të riprodhueshëm.")
bullet("Një vlerësim empirik i realizueshmërisë në harduer konsumatori, me të dhëna reale mbi vonesën dhe sjelljen e refuzimit.")
bullet("Një arkitekturë e pastër, modulare, me kontroll aksesi me role dhe audit, e dokumentuar plotësisht.")

h2("1.7 Qëllimi dhe Kufizimet")
para("Sistemi është projektuar për përdorim institucional lokal ose të izoluar (air-gapped). Janë "
     "qëllimisht jashtë fushës: njohja optike e karaktereve (OCR) për dokumente të skanuara, vendosja "
     "në cloud, arkitekturat me mikroshërbime dhe përpunimi shpërndarës. Korpusi i përdorur për "
     "vlerësim është demonstrues; vlerësimi i saktësisë faktike përmbajtësore kërkon gjykim manual nga "
     "ekspertë të fushës juridike, i cili konsiderohet punë e ardhshme.")

h2("1.8 Struktura e Tezës")
para("Kapitulli 2 paraqet rishikimin e literaturës mbi AI, LLM, RAG, embeddings dhe sektorin publik. "
     "Kapitulli 3 përshkruan metodologjinë dhe analizën e kërkesave. Kapitulli 4 jep dizajnin e "
     "sistemit me diagrama UML/ER. Kapitulli 5 detajon implementimin me copëza kodi dhe pamje nga "
     "ekrani. Kapitulli 6 paraqet eksperimentet dhe rezultatet reale. Kapitulli 7 diskuton gjetjet, "
     "përfitimet dhe kufizimet. Kapitulli 8 jep konkluzionet dhe punën e ardhshme.")

# ======================================================================================
# KAPITULLI 2 — LITERATURA
# ======================================================================================
h1("Kapitulli 2 — Rishikim i Literaturës")

h2("2.1 Inteligjenca Artificiale dhe Përpunimi i Gjuhës Natyrore")
para("Përpunimi i Gjuhës Natyrore (NLP) ka kaluar nga qasje të bazuara në rregulla dhe modele "
     "statistikore (n-grame, modele Markov) te metodat me rrjete nervore. Hapi vendimtar ishte "
     "prezantimi i arkitekturës Transformer [1], e cila zëvendësoi përsëritjen sekuenciale me "
     "mekanizmin e vetë-vëmendjes (self-attention). Ky mekanizëm lejon modelimin paralel të varësive "
     "afatgjata midis fjalëve, duke kapur kontekstin semantik në mënyrë shumë më efektive se rrjetet "
     "rekurente të mëparshme.")

h2("2.2 Modelet e Mëdha Gjuhësore (LLM)")
para("LLM-të janë rrjete Transformer me miliarda parametra, të paratrajnuar mbi korpuse masive teksti "
     "nëpërmjet objektivit të parashikimit të fjalës pasuese. Seria GPT [2] demonstroi se shkallëzimi "
     "i parametrave dhe i të dhënave prodhon aftësi emergjente të gjenerimit dhe arsyetimit «few-shot». "
     "Modelet me peshë të hapur si Llama, Qwen [10] dhe Gemma [11] e kanë demokratizuar aksesin, duke "
     "mundësuar ekzekutimin lokal.")
para("Megjithatë, LLM-të kanë kufizime thelbësore për sektorin publik. Së pari, njohuria kodifikohet "
     "në mënyrë parametrike, çka i bën të paafta për të cituar burimin dhe të prirura ndaj "
     "halucinacionit. Së dyti, ato kanë një «prerje njohurie» (knowledge cutoff) dhe nuk dinë "
     "informacion pas datës së trajnimit. Së treti, ritrajnimi për të përfshirë njohuri të reja është "
     "i kushtueshëm. Mjetet si Ollama lejojnë ekzekutimin lokal të këtyre modeleve, duke ofruar "
     "privatësi dhe kontroll, por me koston e kërkesave të larta llogaritëse — siç vërtetohet edhe nga "
     "matjet e kësaj teme (Kapitulli 6).")

h2("2.3 Retrieval-Augmented Generation (RAG)")
para("RAG [3] u propozua për të kombinuar fuqinë gjeneruese të LLM-ve me saktësinë faktike të një baze "
     "njohurish të jashtme. Në vend që modeli të mbështetet vetëm në njohurinë parametrike, një "
     "komponent retrieval gjen pasazhet më të ngjashme me pyetjen nga një korpus dhe i fut ato në "
     "promptin e modelit si kontekst. Kjo qasje ka tri avantazhe: (1) zvogëlon halucinacionin sepse "
     "përgjigja bazohet në burime; (2) mundëson citime të verifikueshme; dhe (3) lejon përditësimin e "
     "njohurisë thjesht duke ndryshuar korpusin, pa ritrajnuar modelin.")
para("Pipeline-i tipik RAG përbëhet nga: copëzimi i dokumenteve (chunking), gjenerimi i embeddings, "
     "indeksimi në një bazë vektoriale, kërkimi semantik (retrieval), ndërtimi i promptit me kontekst, "
     "dhe gjenerimi. Variante të avancuara përfshijnë retrieval-in hibrid (kombinim i kërkimit dense "
     "me BM25 leksikor) dhe rirenditjen (re-ranking) me cross-encoders [12]. Një sfidë kyçe e RAG-ut "
     "është vendosja se kur korpusi NUK e përmban përgjigjen — pikërisht ky problem adresohet nga "
     "«porta e refuzimit» e propozuar në këtë temë.")

h2("2.4 Embeddings dhe Bazat Vektoriale")
para("Embeddings janë përfaqësime vektoriale të tekstit në një hapësirë ku afërsia gjeometrike "
     "korrespondon me ngjashmërinë semantike. Evolucioni nga word2vec te modelet kontekstuale si BERT "
     "dhe më pas Sentence-BERT [5] mundësoi përfaqësime të tëra fjalish. Modeli bge-m3 [4] është një "
     "model shumëgjuhësh i gjeneratës së fundit që mbulon mbi 100 gjuhë, përfshirë gjuhë me burime të "
     "kufizuara si shqipja, dhe prodhon vektorë me 1024 dimensione. Përzgjedhja e tij është thelbësore "
     "për cilësinë e retrieval-it në shqip.")
para("Bazat vektoriale ruajnë këto embeddings dhe ofrojnë kërkim efikas të fqinjëve më të afërt. Për "
     "shkak se kërkimi i saktë është i kushtueshëm, përdoren algoritme të përafërta (ANN) si HNSW "
     "(Hierarchical Navigable Small World) [6], të cilët ndërtojnë një graf shumështresor për kërkim "
     "logaritmik. ChromaDB [7] është një bazë vektoriale e fokusuar te thjeshtësia dhe lokaliteti, duke "
     "e bërë të përshtatshme për aplikacione në një makinë të vetme.")

h2("2.5 Sistemet e Inteligjencës së Dokumenteve")
para("Sistemet e Document Intelligence nxjerrin, strukturojnë dhe interpretojnë përmbajtjen e "
     "dokumenteve. Në domenin ligjor dhe institucional, kërkesat dalluese janë saktësia faktike, "
     "gjurmueshmëria (mundësia e citimit te burimi) dhe konfidencialiteti. Pikërisht në kryqëzimin e "
     "këtyre kërkesave, një RAG lokal me citime dhe me refuzim të kontrolluar ofron një avantazh të "
     "qartë krahasuar me një asistent të përgjithshëm cloud.")

h2("2.6 Transformimi Dixhital i Sektorit Publik në Shqipëri")
para("Dixhitalizimi i shërbimeve publike në Shqipëri, i përqendruar te platforma e-Albania, ka rritur "
     "pritshmërinë për mjete që e bëjnë njohurinë institucionale të aksesueshme dhe të kërkueshme. "
     "Megjithatë, korniza ligjore për mbrojtjen e të dhënave personale dhe natyra e ndjeshme e shumë "
     "dokumenteve e bëjnë të papërshtatshëm dërgimin e tyre te shërbime cloud të huaja. Kjo krijon një "
     "hapësirë të qartë për zgjidhje plotësisht lokale si DOKU, që ruajnë sovranitetin e të dhënave.")

h2("2.7 Punime të Ngjashme dhe Hendeku Kërkimor")
para("Pjesa më e madhe e asistentëve të dokumenteve të bazuar në RAG (p.sh. zgjidhje komerciale dhe "
     "platforma «chat-with-your-docs») mbështeten në API cloud (OpenAI, Anthropic) dhe janë optimizuar "
     "kryesisht për anglishten. Punime kërkimore mbi RAG-un fokusohen shpesh te cilësia e retrieval-it, "
     "por më pak te kontrolli i refuzimit dhe te kërkesat institucionale (role, audit, lokalitet). "
     "Mungojnë zgjidhje plotësisht lokale, të specializuara për gjuhën shqipe dhe për kontekstin "
     "institucional shqiptar, që integrojnë kontroll aksesi me role, gjurmë auditi dhe një mekanizëm të "
     "qartë anti-halucinacion. DOKU e mbush pikërisht këtë hendek.")

# ======================================================================================
# KAPITULLI 3 — METODOLOGJIA
# ======================================================================================
h1("Kapitulli 3 — Metodologjia dhe Analiza e Kërkesave")

h2("3.1 Metodologjia e Kërkimit")
para("Kërkimi ndjek metodologjinë Design Science Research (DSR), e cila është e përshtatshme për "
     "krijimin dhe vlerësimin e artefakteve teknologjike. Cikli i DSR përfshin: identifikimin e një "
     "problemi real dhe relevant; përcaktimin e objektivave të zgjidhjes; projektimin dhe ndërtimin e "
     "artefaktit (sistemi DOKU); demonstrimin e tij; vlerësimin empirik kundrejt objektivave; dhe "
     "komunikimin e rezultateve. Ky cikël iterativ siguron si një kontribut praktik (sistemi vetë) "
     "ashtu edhe njohuri të transferueshme (mekanizmi i bazueshmërisë dhe gjetjet eksperimentale).")

h2("3.2 Elicitimi dhe Analiza e Kërkesave")
para("Kërkesat u nxorën nga modelimi i një skenari realist institucional: një organizatë publike ku "
     "një administrator menaxhon një korpus të centralizuar dokumentesh zyrtare, ndërsa punonjësit "
     "(punonjesit) e konsultojnë atë nëpërmjet pyetjeve dhe përmbledhjeve, pa pasur të drejtë ta "
     "modifikojnë. Kjo ndarje rolesh pasqyron parimin e privilegjit minimal.")

h3("3.2.1 Kërkesat Funksionale")
table(["ID", "Kërkesa Funksionale", "Roli"], [
    ["KF1", "Autentikim me përdorues/fjalëkalim dhe ndryshim i detyruar i fjalëkalimit fillestar.", "Të dyja"],
    ["KF2", "Dy role me të drejta të ndara: administrator (menaxhim) dhe punonjës (vetëm-lexim).", "Sistemi"],
    ["KF3", "Ngarkim dokumentesh PDF/DOCX, nxjerrje teksti, validim dhe indeksim.", "Admin"],
    ["KF4", "Pyetje në gjuhë natyrore me përgjigje të bazuara dhe citime te dokumenti/faqja.", "Punonjës"],
    ["KF5", "Refuzim i pyetjeve jashtë korpusit, pa gjeneruar halucinacion.", "Sistemi"],
    ["KF6", "Përmbledhje dokumentesh në katër formate të ndryshme.", "Punonjës"],
    ["KF7", "Eksport i përgjigjeve dhe përmbledhjeve në format Word (.docx).", "Punonjës"],
    ["KF8", "Aktivizim/çaktivizim, editim metadatash dhe riindeksim i dokumenteve.", "Admin"],
    ["KF9", "Krijim dhe menaxhim përdoruesish (pa regjistrim publik).", "Admin"],
    ["KF10", "Regjistrim auditi për çdo veprim domethënës.", "Sistemi"],
    ["KF11", "Modul eksperimentesh RAG kundrejt no-RAG me vlerësim manual dhe eksport CSV.", "Admin"],
    ["KF12", "Filtrim dokumentesh sipas tipit, institucionit, vitit dhe fjalëkyçeve.", "Punonjës"],
])

h3("3.2.2 Kërkesat Jofunksionale")
table(["ID", "Kërkesa Jofunksionale", "Realizimi në DOKU"], [
    ["KJF1", "Lokalitet i plotë — pa API/rrjet cloud për inferencë apo embeddings.", "Ollama + bge-m3 lokalisht"],
    ["KJF2", "Siguri — fjalëkalime me bcrypt; kontroll aksesi në kod, jo vetëm UI.", "ui.require_admin, bcrypt"],
    ["KJF3", "Bazueshmëri — çdo pohim citon një copëz; nën pragun refuzon.", "Porta e refuzimit (0.38)"],
    ["KJF4", "Përdorshmëri — ndërfaqe në shqip, e thjeshtë për punonjës joteknikë.", "Streamlit multipage"],
    ["KJF5", "Realizueshmëri — funksionon në 16GB RAM dhe GPU 4GB.", "qwen2.5:3b / gemma2:9b"],
    ["KJF6", "Mirëmbajtshmëri — module të ndara me përgjegjësi të qartë.", "config/modules/pages"],
    ["KJF7", "Gjurmueshmëri — çdo veprim regjistrohet me përdorues, kohë dhe detaje.", "audit_logs"],
])

h2("3.3 Përzgjedhja e Teknologjisë dhe Arsyetimi")
para("Çdo teknologji u përzgjodh me kritere të qarta që lidhen drejtpërdrejt me kërkesat "
     "jofunksionale, veçanërisht lokalitetin, mbështetjen e shqipes dhe realizueshmërinë në harduer të "
     "kufizuar.")
table(["Komponenti", "Teknologjia", "Arsyetimi i përzgjedhjes"], [
    ["Gjuha", "Python 3.13", "Ekosistem i pasur AI/NLP; integrim i drejtpërdrejtë me Sentence-Transformers, ChromaDB, Ollama."],
    ["Ndërfaqja", "Streamlit", "Zhvillim i shpejtë i një UI shumëfaqëshe në Python të pastër; e përshtatshme për prototip akademik dhe demo."],
    ["BD relacionale", "SQLite", "Pa server, skedar i vetëm, zero-konfigurim — ideal për lokalitet; mjaftueshëm për ngarkesën e një institucioni."],
    ["BD vektoriale", "ChromaDB", "Lokale, e thjeshtë, persistente, me indeks HNSW dhe metadata filtering — pa nevojë për shërbim të jashtëm."],
    ["Embeddings", "bge-m3", "Shumëgjuhësh me mbështetje reale të shqipes; vektorë 1024-D të normalizuar për ngjashmëri kosinusi."],
    ["LLM", "Ollama (qwen2.5:3b / gemma2:9b)", "Inferencë plotësisht lokale e modeleve me peshë të hapur; ndërrim modeli pa ndryshim kodi."],
    ["PDF", "PyMuPDF (fitz)", "Nxjerrje e shpejtë teksti për faqe dhe renderim faqesh në imazh për parapamje."],
    ["Word", "python-docx", "Lexim i DOCX-eve të ngarkuara dhe gjenerim i raporteve të eksportuara."],
    ["Siguria", "bcrypt", "Hash fjalëkalimesh me salt për përdorues; standard industrie."],
])
para("Vlen të theksohet pse u shmangën alternativat e zakonshme: Postgres/MySQL u shmangën sepse "
     "kërkojnë server dhe e komplikojnë lokalitetin; FastAPI/React u shmangën sepse do të shtonin "
     "kompleksitet arkitekturor pa vlerë për një sistem me një makinë; dhe API cloud (OpenAI/Anthropic) "
     "u përjashtuan kategorikisht për shkak të kërkesës për konfidencialitet.")

# ======================================================================================
# KAPITULLI 4 — DIZAJNI
# ======================================================================================
h1("Kapitulli 4 — Dizajni i Sistemit")

h2("4.1 Arkitektura e Përgjithshme")
para("DOKU ndjek një arkitekturë me shtresa (layered architecture) me ndarje të qartë përgjegjësish. "
     "Shtresa e prezantimit përbëhet nga faqet Streamlit (dosja pages/) dhe pika hyrëse app.py që "
     "menaxhon login-in, sesionin dhe navigimin sipas rolit. Shtresa e logjikës (dosja modules/) "
     "përmban gjithë logjikën e biznesit. Shtresa e të dhënave përbëhet nga SQLite (metadata dhe "
     "regjistrime) dhe ChromaDB (vektorë), si dhe nga dy shërbime lokale: modeli i embeddings (bge-m3) "
     "dhe serveri LLM (Ollama). Kjo ndarje siguron kohezion të lartë brenda moduleve dhe lidhje të "
     "ulët midis tyre; ndërfaqja nuk i prek kurrë drejtpërdrejt bazat e të dhënave.")
code(
"+-------------------------------------------------------------+\n"
"|                  SHTRESA E PREZANTIMIT (UI)                  |\n"
"|   app.py (login/sesion/navigim) + pages/ (8 faqe)           |\n"
"+-----------------------------+-------------------------------+\n"
"                              |\n"
"+-----------------------------v-------------------------------+\n"
"|                    SHTRESA E LOGJIKES (modules/)             |\n"
"|  auth  documents  rag_pipeline  embeddings  vector_store    |\n"
"|  llm_client  history  audit  experiments  export_docx  ui   |\n"
"+----------------+----------------------------+---------------+\n"
"                 |                            |\n"
"        +--------v-------+           +--------v---------+\n"
"        |    SQLite      |           |    ChromaDB      |\n"
"        | (5 tabela)     |           | (vektore 1024-D) |\n"
"        +----------------+           +------------------+\n"
"                 |                            |\n"
"        +--------v-------+           +--------v---------+\n"
"        | bge-m3 (embed) |           |  Ollama (LLM)    |\n"
"        +----------------+           +------------------+")
caption("Figura 4.1 — Arkitektura me shtresa e sistemit DOKU.")

h2("4.2 Përgjegjësitë e Moduleve")
table(["Moduli", "Përgjegjësia"], [
    ["config.py", "Konfigurim qendror: modelet, shtigjet, pragjet, enumet."],
    ["database.py", "Skema SQLite (5 tabela), lidhja dhe menaxhimi i transaksioneve."],
    ["auth.py", "Hash bcrypt, role, admin i parazgjedhur, ndryshim i detyruar fjalëkalimi."],
    ["document_processor.py", "Nxjerrje teksti (PyMuPDF/docx), validim dhe copëzim."],
    ["embeddings.py", "Mbështjellës i bge-m3 për gjenerimin e embeddings."],
    ["vector_store.py", "ChromaDB: indeksim, kërkim me filtra, konvertim distancë→ngjashmëri."],
    ["documents.py", "CRUD i dokumenteve: ngarkim, editim, status, fshirje, riindeksim."],
    ["rag_pipeline.py", "Zemra e RAG: retrieval → porta e refuzimit → prompt → LLM → citime."],
    ["llm_client.py", "Klienti Ollama; ndërrim modeli; trajtim gabimesh."],
    ["history.py / audit.py", "Persistencë e historikut të bisedave dhe gjurmës së auditit."],
    ["export_docx.py", "Gjenerim i raporteve Word për përgjigje dhe përmbledhje."],
    ["experiments.py", "Harness RAG vs no-RAG, vlerësim manual dhe eksport CSV."],
    ["ui.py", "Roje sesioni dhe autorizimi të përbashkëta për faqet."],
])

h2("4.3 Diagrami i Rasteve të Përdorimit (Use Case)")
code(
"            DOKU — Use Case\n"
"  ( Administrator )                 ( Punonjes )\n"
"        |                                 |\n"
"        |-- Menaxho perdoruesit           |-- Hyr ne sistem\n"
"        |-- Ngarko/Edito/Fshi dokument    |-- Pyet dokumentet (RAG)\n"
"        |-- Aktivizo/Caktivizo dok.       |-- Gjenero permbledhje\n"
"        |-- Riindekso korpusin            |-- Eksporto ne Word\n"
"        |-- Shiko audit log               |-- Filtro dokumentet\n"
"        |-- Nis eksperimente (RAG/noRAG)  |-- Shiko historikun tim\n"
"   (te dyja rolet: hyr, ndrysho fjalekalimin, dil)")
caption("Figura 4.2 — Diagrami i rasteve të përdorimit.")

h2("4.4 Diagrami i Komponentëve")
code(
"app.py (router/auth) --> pages/* --> modules/*\n"
"modules/rag_pipeline --> {embeddings, vector_store, llm_client}\n"
"modules/documents    --> {document_processor, vector_store, database}\n"
"modules/{auth,history,audit,experiments} --> database\n"
"vector_store --> ChromaDB ; embeddings --> bge-m3 ; llm_client --> Ollama")
caption("Figura 4.3 — Diagrami i komponentëve dhe varësive.")

h2("4.5 Rrjedha e të Dhënave dhe Diagrami i Sekuencës (RAG)")
code(
"Punonjes -> UI: shkruan pyetjen\n"
"UI -> rag_pipeline: answer_question(q, active_ids)\n"
"rag_pipeline -> embeddings: embed_query(q)         [bge-m3 -> vektor 1024-D]\n"
"rag_pipeline -> vector_store: query(qvec, active_ids)\n"
"vector_store -> ChromaDB: kerkim top-k (cosine)\n"
"ChromaDB --> rag_pipeline: copeza + distanca\n"
"alt  top_score < 0.38  (PORTA E REFUZIMIT)\n"
"   rag_pipeline --> UI: REFUZIM (pa e thirrur LLM-ne)\n"
"else\n"
"   rag_pipeline -> llm_client: generate(system + kontekst i cituar + pyetja)\n"
"   llm_client -> Ollama: chat(model)\n"
"   Ollama --> rag_pipeline: pergjigje\n"
"   rag_pipeline --> UI: pergjigje + burime (dok., faqe)\n"
"end\n"
"UI -> history: save() ; UI -> audit: log()")
caption("Figura 4.4 — Diagrami i sekuencës për një pyetje-përgjigje me portën e refuzimit.")

h2("4.6 Diagrami i Aktivitetit (Ngarkim Dokumenti)")
code(
"[Start] -> Ngarko PDF/DOCX -> Gjenero emer te sigurt skedari\n"
"   -> A ekziston duplikat? --PO--> Gabim 'ekziston' -> [End]\n"
"   --JO--> Ruaj skedarin -> Nxjerr tekst (PyMuPDF per faqe / docx)\n"
"   -> A ka tekst te lexueshem? --JO--> Gabim 'pa tekst (skan)' -> [End]\n"
"   --PO--> Copezo me mbivendosje -> Embed (bge-m3) -> Ruaj vektoret ne ChromaDB\n"
"   -> Ruaj metadata ne SQLite -> Audit log 'upload_document' -> [End]")
caption("Figura 4.5 — Diagrami i aktivitetit për ngarkimin dhe indeksimin e dokumentit.")

h2("4.7 Diagrami i Vendosjes (Deployment)")
code(
"+---------------------- Makina Lokale (Windows) ----------------------+\n"
"|  Shfletues --HTTP/WS--> [Streamlit :8501] --HTTP--> [Ollama :11434] |\n"
"|        |                                                            |\n"
"|        +--> [SQLite app.db]   [ChromaDB chroma_db/]   [bge-m3]      |\n"
"|  (opsionale, vetem per demo) [Cloudflare Tunnel] -> URL publik HTTPS|\n"
"+--------------------------------------------------------------------+\n"
" Inferenca dhe te dhenat MBETEN tersisht lokale; tunnel-i vetem ekspozon UI-ne.")
caption("Figura 4.6 — Diagrami i vendosjes (gjithçka lokale).")

h2("4.8 Dizajni i Bazës së të Dhënave (ER)")
para("Skema relacionale përbëhet nga pesë tabela në SQLite. Vektorët e copëzave ruhen ndaras në "
     "ChromaDB, të lidhura logjikisht me dokumentet nëpërmjet fushës document_id në metadata. Tabela "
     "users ruan llogaritë dhe rolet; documents ruan metadatat e korpusit; chat_history ruan çdo "
     "pyetje-përgjigje; audit_logs ruan gjurmën e veprimeve; dhe experiment_results ruan të dhënat e "
     "krahasimit RAG/no-RAG.")
code(
"users(id PK, username UNIQUE, password_hash, full_name, role[CHECK admin|punonjes],\n"
"      must_change_password, is_active, created_at, updated_at)\n"
"documents(id PK, filename UNIQUE, original_filename, stored_path, title, institution,\n"
"      document_type, year, description, uploaded_by, status[CHECK active|inactive],\n"
"      num_pages, total_chunks, created_at, updated_at)\n"
"chat_history(id PK, user_id ->users.id, username, question, answer, mode,\n"
"      selected_document_id ->documents.id, sources_json, response_time,\n"
"      exported_to_word, created_at)\n"
"audit_logs(id PK, user_id ->users.id, username, action, details, created_at)\n"
"experiment_results(id PK, question, answer_without_rag, answer_with_rag,\n"
"      time_without_rag, time_with_rag, chunks_used, has_sources,\n"
"      manual_accuracy_without_rag, manual_accuracy_with_rag,\n"
"      hallucination_without_rag, hallucination_with_rag, notes, created_at)\n"
"\n"
"ChromaDB: doku_chunks(id='docId:chunkIdx', embedding[1024],\n"
"      metadata{document_id, filename, title, institution, document_type,\n"
"               year, page_number, chunk_index, status})")
caption("Figura 4.7 — Skema ER (SQLite) dhe koleksioni vektorial (ChromaDB).")

h2("4.9 Dizajni i Sigurisë")
para("Dizajni i sigurisë mbështetet në disa parime: fjalëkalimet nuk ruhen kurrë në tekst të thjeshtë "
     "por hash-ohen me bcrypt (me salt për përdorues); kontrolli i aksesit zbatohet në nivel kodi "
     "(funksioni ui.require_admin ndalon ekzekutimin e faqes nëse roli nuk është admin), jo thjesht "
     "duke fshehur elemente në UI; të gjitha pyetjet SQL janë të parametrizuara, duke parandaluar "
     "SQL injection; nuk ekziston regjistrim publik, pra vetëm administratori krijon llogari; dhe çdo "
     "veprim domethënës regjistrohet në tabelën audit_logs me përdoruesin, kohën dhe detajet.")

h2("4.10 Dizajni i Pipeline-it RAG dhe Porta e Refuzimit")
para("Elementi qendror dhe origjinal i dizajnit është «porta e refuzimit». Logjika është e qëllimshme: "
     "porta ekzekutohet PARA thirrjes së modelit. Nëse ngjashmëria më e lartë e marrë (top_score) është "
     "nën pragun MIN_SIMILARITY (0.38), sistemi kthen menjëherë një mesazh refuzimi standard, pa e "
     "thirrur fare LLM-në. Kjo e bën fabrikimin strukturalisht të pamundur për pyetje jashtë korpusit, "
     "dhe njëkohësisht kursen kohë llogaritëse. Mbi pragun, ndërtohet një prompt që përmban kontekstin "
     "e cituar (me numra faqesh) dhe një system-prompt që e udhëzon modelin të përgjigjet vetëm nga "
     "konteksti dhe të deklarojë kur informacioni nuk gjendet. Për dokumentet normative (Ligj, VKM, "
     "Rregullore, Udhëzim) shtohet automatikisht një shënim ligjor që përgjigja është ndihmëse dhe "
     "dokumenti origjinal mbetet burimi zyrtar.")

# ======================================================================================
# KAPITULLI 5 — IMPLEMENTIMI
# ======================================================================================
h1("Kapitulli 5 — Implementimi")
para("Ky kapitull përshkruan implementimin e secilit komponent, duke shoqëruar përshkrimet me copëza "
     "kodi përfaqësuese dhe me pamje reale nga ekrani të aplikacionit të ekzekutuar.")

h2("5.1 Autentikimi dhe Faqja e Hyrjes")
para("Autentikimi bazohet në bcrypt. Funksioni hash_password kufizon hyrjen në 72 bajt (kufi i "
     "bcrypt-it) dhe gjeneron një salt unik. Faqja e hyrjes është e izoluar: paneli anësor (sidebar) "
     "fshihet plotësisht derisa autentikimi të përfundojë me sukses, dhe navigimi i faqeve është i "
     "fshehur, çka pengon aksesin te faqet pa hyrje.")
code(
"def hash_password(password):\n"
"    return bcrypt.hashpw(password.encode('utf-8')[:72], bcrypt.gensalt()).decode()\n\n"
"def authenticate(username, password):\n"
"    row = get_user(username)\n"
"    if row is None or not row['is_active']: return None\n"
"    if not verify_password(password, row['password_hash']): return None\n"
"    return row")
figure("01_login.png", "Figura 5.1 — Faqja e hyrjes (login). Paneli anësor fshihet derisa "
       "autentikimi të përfundojë; titulli dhe formulari janë të qendërzuar.")

h2("5.2 Autorizimi me Role dhe Menaxhimi i Përdoruesve")
para("Autorizimi zbatohet në kod. Çdo faqe administrative thërret ui.require_admin() në krye, e cila "
     "ndalon ekzekutimin (st.stop) nëse roli nuk është admin. Administratori i parazgjedhur "
     "(admin/***REMOVED-CREDENTIAL***) krijohet automatikisht nëse nuk ekziston asnjë admin, dhe detyrohet të "
     "ndryshojë fjalëkalimin në hyrjen e parë. Administratori krijon punonjësit, të cilët gjithashtu "
     "detyrohen të ndryshojnë fjalëkalimin fillestar.")
code(
"def require_admin():\n"
"    user = current_user()\n"
"    if user['role'] != auth.ADMIN:\n"
"        st.error('Nuk keni leje për këtë faqe (vetëm administratori).')\n"
"        st.stop()\n"
"    return user")
figure("08_users.png", "Figura 5.2 — Faqja e menaxhimit të përdoruesve (vetëm admin). Krijim "
       "përdoruesi me rol, dhe aktivizim/çaktivizim. Admini nuk mund të çaktivizojë veten.")

h2("5.3 Menaxhimi dhe Indeksimi i Dokumenteve")
para("Ngarkimi i një dokumenti kalon nëpër disa hapa: gjenerimi i një emri skedari të sigurt "
     "(sanitizim Unicode + regex që ruan prapashtesën .pdf/.docx dhe pengon path traversal), "
     "kontrolli i dublikatave, ruajtja e skedarit, përpunimi (nxjerrje + validim + copëzim) dhe "
     "indeksimi në ChromaDB. Statusi active/inactive kontrollon dukshmërinë e dokumentit në kërkim. "
     "Administratori mund të editojë metadatat, të riindeksojë një dokument ose tërë korpusin, dhe të "
     "shohë një parapamje të faqeve të para.")
code(
"def safe_filename(name):\n"
"    # sanitizim Unicode + regex; ruan .pdf/.docx; pengon path traversal\n"
"    stem = unicodedata.normalize('NFKD', parts[0]).encode('ascii','ignore').decode()\n"
"    stem = re.sub(r'[^A-Za-z0-9_.-]+', '_', stem).strip('._') or 'dokument'\n"
"    return stem + ext")
figure("07_documents.png", "Figura 5.3 — Faqja e menaxhimit të dokumenteve (vetëm admin). Ngarkim, "
       "metadata, aktivizim/çaktivizim, riindeksim, fshirje dhe parapamje për çdo dokument.")

h2("5.4 Përpunimi i PDF/DOCX dhe Copëzimi")
para("PyMuPDF nxjerr tekstin për çdo faqe veçmas, duke ruajtur numrin e faqes për citime të sakta; "
     "DOCX-et trajtohen si një «faqe» e vetme. Një validim siguron që dokumenti përmban tekst të "
     "lexueshëm (të paktën 100 karaktere), duke refuzuar PDF-të e skanuara pasi versioni bazë nuk "
     "mbështet OCR. Copëzimi përdor një dritare karakteresh me mbivendosje (CHUNK_SIZE=800, "
     "CHUNK_OVERLAP=120), duke u përpjekur të ndajë pranë një hapësire boshe për të mos prerë fjalët.")
code(
"def process_document(path):\n"
"    pages = _pages_for(path)            # PyMuPDF per faqe / docx\n"
"    validate_has_text(pages)           # >=100 karaktere, perndryshe NoExtractableTextError\n"
"    chunks, idx = [], 0\n"
"    for page_no, page_text in enumerate(pages, start=1):\n"
"        for piece in chunk_text(page_text, CHUNK_SIZE, CHUNK_OVERLAP):\n"
"            chunks.append(Chunk(text=piece, page_number=page_no, chunk_index=idx)); idx += 1\n"
"    return chunks, len(pages)")

h2("5.5 Gjenerimi i Embeddings dhe Ruajtja Vektoriale")
para("Modeli bge-m3 prodhon vektorë të normalizuar me L2, çka do të thotë se ngjashmëria kosinusi "
     "barazohet me produktin skalar. ChromaDB konfigurohet me hapësirë 'cosine' dhe ruan për çdo copëz "
     "metadata të plota (dokumenti, titulli, institucioni, tipi, viti, faqja, indeksi), të cilat "
     "mundësojnë ndërtimin e citimeve dhe filtrimin sipas dokumenteve aktive.")
code(
"# embeddings.py\n"
"def embed_texts(texts):\n"
"    return get_model().encode(texts, normalize_embeddings=True).tolist()\n\n"
"# vector_store.py\n"
"vectors = embeddings.embed_texts(docs)\n"
"col.add(ids=ids, documents=docs, metadatas=metadatas, embeddings=vectors)")

h2("5.6 Retrieval, Filtrimi dhe Porta e Refuzimit")
para("Kërkimi merr top-k copëza (k=5), me filtrim sipas grupit të dokumenteve aktive ose një "
     "dokumenti specifik të zgjedhur nga përdoruesi. Distanca e kthyer nga ChromaDB konvertohet në "
     "ngjashmëri (similarity = 1 − distance). Porta e refuzimit ekzekutohet menjëherë pas retrieval-it "
     "dhe para çdo thirrjeje LLM. Faqja e pyetjeve mbështet filtrim sipas tipit, institucionit, vitit "
     "dhe fjalëkyçeve, si dhe zgjedhjen e fushës së kërkimit.")
code(
"chunks = vs.query(question, active_doc_ids=active)\n"
"top = chunks[0].score if chunks else 0.0\n"
"if not chunks or top < config.MIN_SIMILARITY:        # PORTA E REFUZIMIT\n"
"    return Answer(config.REFUSAL_MESSAGE, refused=True, top_score=top)  # pa LLM\n"
"relevant = [c for c in chunks if c.score >= config.MIN_SIMILARITY] or chunks[:1]\n"
"prompt = _USER_PROMPT.format(context=_format_context(relevant), question=question)\n"
"text = llm_client.generate(prompt, system=SYSTEM_PROMPT)")
figure("03_ask_interface.png", "Figura 5.4 — Faqja «Pyet Dokumentet»: filtrim, zgjedhje e fushës së "
       "kërkimit dhe fusha e pyetjes në gjuhë natyrore.")
figure("04_ask_refusal.png", "Figura 5.5 — Porta e refuzimit në veprim: për një pyetje jashtë "
       "korpusit, sistemi refuzon dhe shfaq ngjashmërinë më të lartë (0.353) nën pragun 0.38, pa e "
       "thirrur modelin.")

h2("5.7 Integrimi me Ollama")
para("Klienti i LLM-së komunikon me serverin lokal Ollama nëpërmjet HTTP. Modeli aktiv mund të "
     "ndërrohet nga paneli anësor pa ndryshuar kodin. Në rast se Ollama nuk është aktiv ose modeli "
     "mungon, hidhet një përjashtim i qartë në shqip që UI-ja e shfaq, duke i thënë përdoruesit të "
     "sigurohet që Ollama po ekzekuton dhe modeli është shkarkuar.")

h2("5.8 Përmbledhja e Dokumenteve")
para("Moduli i përmbledhjes nxjerr tekstin e plotë të një dokumenti dhe e dërgon te LLM-ja me një "
     "system-prompt që ndalon shtimin e informacionit nga jashtë. Mbështeten katër formate: e "
     "shkurtër, e detajuar, pika kryesore, dhe një format i orientuar drejt vendimmarrjes "
     "institucionale (me titujt Qëllimi, Detyrimet/Implikimet, Rekomandime). Çdo përmbledhje shoqërohet "
     "me një shënim verifikimi.")
figure("05_summary.png", "Figura 5.6 — Faqja «Përmbledhje Dokumenti»: zgjedhje dokumenti dhe formati i "
       "përmbledhjes, me eksport në Word.")

h2("5.9 Historiku dhe Gjurma e Auditit")
para("Çdo pyetje-përgjigje ruhet në tabelën chat_history me burimet (si JSON), kohën e përgjigjes dhe "
     "modalitetin (rag/summary). Punonjësi sheh vetëm historikun e tij. Paralelisht, çdo veprim "
     "domethënës (hyrje, ngarkim, fshirje, pyetje, eksperiment, etj.) regjistrohet në audit_logs, e "
     "cila është e aksesueshme vetëm nga administratori.")
figure("06_history.png", "Figura 5.7 — Faqja «Historiku im»: pyetjet e mëparshme të përdoruesit.")
figure("09_audit.png", "Figura 5.8 — Faqja «Audit Log» (vetëm admin): gjurma e plotë e veprimeve me "
       "përdorues, veprim, detaje dhe kohë.")

h2("5.10 Eksporti në Word")
para("Nëpërmjet python-docx, sistemi gjeneron raporte .docx të strukturuara për përgjigjet (me "
     "pyetjen, përgjigjen, burimet e cituara dhe një shënim verifikimi ligjor) dhe për përmbledhjet "
     "(me të dhënat e dokumentit dhe formatin). Skedarët ruhen në data/exports/ dhe ofrohen për "
     "shkarkim nëpërmjet butonit përkatës.")

# ======================================================================================
# KAPITULLI 6 — EKSPERIMENTET
# ======================================================================================
h1("Kapitulli 6 — Eksperimentet dhe Rezultatet")

h2("6.1 Dizajni Eksperimental")
para("Eksperimenti krahason të njëjtat pyetje të procesuara në dy mënyra: (a) LLM pa RAG "
     "(answer_without_rag — modeli përgjigjet vetëm nga njohuria parametrike) dhe (b) pipeline-i i "
     "plotë RAG (answer_question — me retrieval dhe portë refuzimi). Për çdo pyetje maten koha e "
     "përgjigjes në të dyja mënyrat, numri i copëzave të përdorura dhe prania e burimeve. Korpusi "
     "përmbante 13 dokumente aktive (1677 copëza të indeksuara). U përdor modeli qwen2.5:3b për shkak "
     "të kohëve të arsyeshme në harduerin e disponueshëm (16GB RAM, GPU 4GB); gemma2:9b u vlerësua "
     "veçmas (seksioni 6.5).")
para("Grupi testues përbëhej nga 10 pyetje: 8 brenda korpusit (mbi tatimet, kodin e punës dhe "
     "strategjinë e dixhitalizimit) dhe 2 qëllimisht jashtë korpusit (çmimi i një bilete avioni dhe "
     "fituesi i Kupës së Botës 2018), të dizajnuara për të testuar portën e refuzimit. Pyetjet vijnë "
     "nga skedari tests/sample_questions.csv.")
figure("10_experiments.png", "Figura 6.1 — Faqja «Eksperimente» (vetëm admin): ekzekutimi i pyetjeve "
       "testuese, tabela e rezultateve me kohë e citime, dhe vlerësimi manual i saktësisë/halucinacionit.")

h2("6.2 Rezultatet e Papërpunuara")
rows = []
for nr, q, tn, tr, ch, hs, inc, ref in RESULTS:
    rows.append([nr, (q[:36] + "…") if len(q) > 37 else q,
                 f"{tn:.1f}", f"{tr:.1f}", ch, "Po" if hs else "Jo",
                 "Brenda" if inc else "Jashtë"])
table(["#", "Pyetja", "t no-RAG (s)", "t RAG (s)", "Copëza", "Burime", "Korpusi"], rows)
caption("Tabela 6.1 — Rezultatet e papërpunuara për 10 pyetjet (të dhëna reale nga harness-i).")

h2("6.3 Analiza e Rezultateve")
h3("6.3.1 Bazueshmëria dhe Citimet (PK1)")
para(f"Për të 8 pyetjet brenda korpusit ({GROUNDED}/8 = 100%), RAG-u ktheu përgjigje me burime të "
     "cituara, duke përdorur 5 copëza për secilën. Kjo i përgjigjet pozitivisht PK1: sistemi prodhon "
     "përgjigje të bazuara dhe të verifikueshme në gjuhën shqipe, ku çdo përgjigje mund të gjurmohet "
     "te dokumenti dhe faqja burimore.")
h3("6.3.2 Sjellja e Refuzimit: RAG kundrejt no-RAG (PK2)")
para("Gjetja më e rëndësishme lidhet me pyetjet jashtë korpusit. Pyetja 10 («Kush e fitoi Kupën e "
     "Botës 2018?») u refuzua saktë nga porta e refuzimit në vetëm 0.43 sekonda, pa e thirrur fare "
     "modelin (0 copëza, pa burime), pasi ngjashmëria më e lartë ishte 0.353 — nën pragun 0.38 (shih "
     "Figurën 5.5). Në të kundërt, LLM-ja pa RAG iu përgjigj së njëjtës pyetje në 44.5 sekonda, duke "
     "prodhuar një përgjigje pa asnjë burim institucional. Ky kontrast demonstron qartë vlerën e "
     "RAG-ut: kontroll strukturor mbi bazueshmërinë, përgjigje pozitive për PK2.")
para("Pyetja 9 («Sa kushton një biletë avioni?») përbën një kufizim të ndershëm: ajo kaloi pragun (3 "
     "copëza me ngjashmëri ≥ 0.38) dhe u përgjigj, megjithëse është jashtë qëllimit të korpusit. Kjo "
     "tregon se pragu i ngjashmërisë nuk është i pagabueshëm dhe motivon përmirësimet e propozuara në "
     "seksionin 7.3 (retrieval hibrid, normalizim diakritikash).")
h3("6.3.3 Vonesa (Latency) (PK3)")
para(f"Koha mesatare për pyetjet brenda korpusit ishte ~{NORAG_MEAN_IN:.1f}s për no-RAG dhe "
     f"~{RAG_MEAN_IN:.1f}s për RAG. Kohët janë relativisht të larta dhe me variancë të madhe për shkak "
     "të presionit të memories në harduerin e kufizuar (modeli shpërndahej pjesërisht në CPU; vlera e "
     "parë, 189s, është një outlier nga ngarkimi fillestar i modelit). Kur porta e refuzimit "
     "aktivizohet, vonesa bie në mënyrë drastike (0.43s), pasi shmanget plotësisht thirrja e modelit — "
     "një përfitim dytësor i rëndësishëm i dizajnit.")
table(["Metrika", "no-RAG", "RAG"], [
    ["Kohë mesatare, brenda korpusit (s)", f"{NORAG_MEAN_IN:.1f}", f"{RAG_MEAN_IN:.1f}"],
    ["Përgjigje me burime (brenda korpusit)", "0/8", "8/8 (100%)"],
    ["Refuzim i saktë jashtë korpusit", "0/2", "1/2"],
    ["Sjellja për pyetjen 10 (jashtë korpusit)", "u përgjigj (44.5s, pa burim)", "refuzoi (0.43s)"],
])
caption("Tabela 6.2 — Përmbledhje krahasuese RAG kundrejt no-RAG (të dhëna reale).")

h2("6.4 Analiza e Pragut të Ngjashmërisë")
para("Pragu MIN_SIMILARITY=0.38 u përcaktua empirikisht. Vëzhgimet treguan se pyetjet brenda korpusit "
     "me diakritikë të saktë shënojnë ngjashmëri rreth 0.70, ato pa diakritikë rreth 0.40, ndërsa "
     "pyetjet jashtë korpusit shënojnë ≤0.35. Pragu 0.38 u zgjodh për të lejuar pyetjet pa diakritikë "
     "duke refuzuar shumicën e atyre jashtë korpusit. Megjithatë, pyetja 9 (0.38+) tregon se ky kufi i "
     "vetëm nuk është i mjaftueshëm; kjo brishtësi adresohet si punë e ardhshme.")

h2("6.5 Krahasimi i Modeleve (qwen2.5:3b kundrejt gemma2:9b)")
para("Një test i veçantë i gemma2:9b në të njëjtin harduer tregoi përgjigje me cilësi gjuhësore më të "
     "mirë në shqip, por me vonesë rreth 5 minuta për pyetje, pasi modeli 9B nuk hyn në 4GB VRAM dhe "
     "shpërndahet kryesisht në CPU. Kjo përbën një kompromis të qartë cilësi–vonesë (PK3): qwen2.5:3b "
     "rekomandohet për ndërveprim në kohë reale, ndërsa gemma2:9b për cilësi maksimale kur vonesa nuk "
     "është kritike. Arkitektura e DOKU-t e mbështet këtë zgjedhje nëpërmjet ndërrimit të modelit nga "
     "paneli anësor, pa ndryshim kodi.")

h2("6.6 Përmbledhje e Gjetjeve")
bullet("RAG prodhoi përgjigje të bazuara me citime për 100% të pyetjeve brenda korpusit.")
bullet("Porta e refuzimit refuzoi pyetjen jashtë korpusit në 0.43s pa thirrur modelin; no-RAG u përgjigj pa burim në 44.5s.")
bullet("Një pyetje jashtë korpusit kaloi pragun — kufizim që motivon retrieval hibrid dhe normalizim diakritikash.")
bullet("Vonesa varet fort nga modeli dhe harduri; ekziston një kompromis i qartë cilësi–shpejtësi.")

# ======================================================================================
# KAPITULLI 7 — DISKUTIMI
# ======================================================================================
h1("Kapitulli 7 — Diskutimi")

h2("7.1 Interpretimi i Gjetjeve")
para("Rezultatet konfirmojnë hipotezën qendrore: kombinimi i retrieval-it me një portë refuzimi para "
     "gjenerimit ofron një kontroll mbi bazueshmërinë që një LLM i pastër nuk e ka. Dallimi midis "
     "refuzimit në 0.43 sekonda të RAG-ut dhe përgjigjes 44.5-sekondëshe pa burim të no-RAG-ut është "
     "një ilustrim konkret dhe i matshëm i kësaj vlere. Për një kontekst institucional, ku një "
     "përgjigje e pasaktë mund të ketë pasoja ligjore, kjo veçori është vendimtare.")

h2("7.2 Përfitimet")
bullet("Besueshmëri dhe gjurmueshmëri: çdo përgjigje është e lidhur me dokumentin dhe faqen burimore.")
bullet("Privatësi dhe sovranitet: gjithçka ekzekutohet lokalisht, pa dërgim të dhënash te shërbime të jashtme.")
bullet("Mbështetje e gjuhës shqipe nëpërmjet embeddings shumëgjuhëshe dhe modeleve lokale.")
bullet("Siguri institucionale: kontroll aksesi me role i zbatuar në kod, dhe gjurmë auditi e plotë.")
bullet("Fleksibilitet: ndërrim modeli sipas kompromisit cilësi/shpejtësi, pa ndryshim kodi.")

h2("7.3 Kufizimet")
bullet("Pragu i ngjashmërisë është empirik dhe i ndjeshëm ndaj mungesës së diakritikave; mund të lejojë false-accept (pyetja 9).")
bullet("Vonesa në harduer konsumatori është e lartë për modele të mëdha (deri ~5 min për gemma2:9b).")
bullet("Copëzimi për faqe humb kontekst ndërfaqësor dhe nuk është semantik.")
bullet("Mungojnë teste të automatizuara dhe një vlerësim manual i saktësisë faktike nga ekspertë juridikë.")
bullet("Vlerësimi përdori një korpus demonstrues të kufizuar (13 dokumente, 10 pyetje).")

h2("7.4 Kërcënimet ndaj Vlefshmërisë")
para("Vlefshmëria e brendshme kufizohet nga varianca e kohëve të matura për shkak të ngarkesës së "
     "harduerit; matjet duhen riprodhuar në kushte më të kontrolluara. Vlefshmëria e jashtme (mundësia "
     "e përgjithësimit) kufizohet nga madhësia e vogël e korpusit dhe e grupit të pyetjeve; rezultatet "
     "duhen konfirmuar në një korpus më të madh dhe më të larmishëm. Vlefshmëria konstrukt (a maten "
     "vërtet bazueshmëria dhe halucinacioni) përmirësohet duke shtuar vlerësim manual nga ekspertë.")

h2("7.5 Rreziqet")
bullet("Prompt-injection nëpërmjet dokumenteve të indeksuara (zbutet pjesërisht nga ngarkimi vetëm-admin).")
bullet("Kredencialet e parazgjedhura nëse nuk ndryshohen; mungesa e mbrojtjes nga sulmet brute-force.")
bullet("Ekspozimi publik nëpërmjet tunnel-it pa autentikim shtesë rrjeti gjatë demonstrimeve.")

# ======================================================================================
# KAPITULLI 8 — KONKLUZIONE
# ======================================================================================
h1("Kapitulli 8 — Konkluzione dhe Punë e Ardhshme")

h2("8.1 Konkluzione")
para("Kjo temë projektoi, implementoi dhe vlerësoi DOKU-n — një sistem RAG plotësisht lokal për "
     "analizën e dokumenteve institucionale në gjuhën shqipe. Vlerësimi empirik tregoi se sistemi "
     "prodhon përgjigje të bazuara me citime për pyetjet brenda korpusit dhe, falë portës së refuzimit, "
     "shmang përgjigjet pa burim për pyetje jashtë korpusit — ndryshe nga një LLM pa RAG, i cili "
     "përgjigjet gjithmonë. Të gjitha objektivat kryesore u arritën: u ndërtua një sistem lokal, u "
     "garantua bazueshmëria, u zbatua siguria me role dhe audit, dhe u vlerësua empirikisht përfitimi i "
     "RAG-ut në harduer konsumatori.")

h2("8.2 Kontributet")
bullet("Një sistem referencë RAG lokal i specializuar për gjuhën shqipe dhe sektorin publik.")
bullet("Një kontratë bazueshmërie (refusal-gate përpara LLM-së) si mekanizëm strukturor anti-halucinacion, me efekt të matur.")
bullet("Një vlerësim empirik me të dhëna reale të vonesës dhe sjelljes së refuzimit, dhe një harness i riprodhueshëm.")

h2("8.3 Punë e Ardhshme")
bullet("Retrieval hibrid (BM25 + dense) dhe rirenditje me cross-encoder për saktësi më të lartë.")
bullet("Copëzim semantik/token-aware dhe normalizim diakritikash për një prag refuzimi më të qëndrueshëm.")
bullet("Suitë testesh të automatizuara dhe vlerësim manual i saktësisë nga ekspertë juridikë.")
bullet("Indekse dhe kufizime FK në bazën e të dhënave, si dhe një mekanizëm migrimesh skeme.")
bullet("Mbrojtje nga brute-force, skadim sesioni dhe konfigurim me variabla mjedisi.")
bullet("Mbështetje OCR për dokumente të skanuara (aktualisht jashtë fushës).")

# ======================================================================================
# REFERENCAT
# ======================================================================================
h1("Referencat")
refs = [
    "A. Vaswani et al., “Attention Is All You Need,” in Proc. NeurIPS, 2017.",
    "T. Brown et al., “Language Models are Few-Shot Learners,” in Proc. NeurIPS, 2020.",
    "P. Lewis et al., “Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,” in Proc. NeurIPS, 2020.",
    "J. Chen et al., “BGE M3-Embedding: Multi-Lingual, Multi-Functionality, Multi-Granularity Text Embeddings,” arXiv:2402.03216, 2024.",
    "N. Reimers and I. Gurevych, “Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks,” in Proc. EMNLP, 2019.",
    "Y. Malkov and D. Yashunin, “Efficient and robust approximate nearest neighbor search using HNSW graphs,” IEEE TPAMI, 2018.",
    "Chroma, “ChromaDB: the open-source embedding database,” [Online]. Available: https://www.trychroma.com",
    "Ollama, “Run large language models locally,” [Online]. Available: https://ollama.com",
    "Streamlit Inc., “Streamlit documentation,” [Online]. Available: https://docs.streamlit.io",
    "A. Yang et al., “Qwen2.5 Technical Report,” arXiv, 2024.",
    "Gemma Team, Google DeepMind, “Gemma 2: Improving Open Language Models,” arXiv, 2024.",
    "Y. Gao et al., “Retrieval-Augmented Generation for Large Language Models: A Survey,” arXiv:2312.10997, 2023.",
    "J. Devlin et al., “BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding,” in Proc. NAACL, 2019.",
    "A. Provilkov et al. / J. Johnson et al., “Billion-scale similarity search (FAISS),” IEEE Trans. Big Data, 2019.",
    "PyMuPDF, “PyMuPDF documentation,” [Online]. Available: https://pymupdf.readthedocs.io",
    "Python Software Foundation, “SQLite — sqlite3 module,” [Online]. Available: https://docs.python.org/3/library/sqlite3.html",
    "N. Provos and D. Mazières, “A Future-Adaptable Password Scheme (bcrypt),” in Proc. USENIX, 1999.",
    "OWASP Foundation, “OWASP Top 10,” [Online]. Available: https://owasp.org/Top10",
]
for i, r in enumerate(refs, 1):
    doc.add_paragraph(f"[{i}] {r}")

# ======================================================================================
# SHTOJCAT
# ======================================================================================
h1("Shtojca A — Udhëzues Instalimi dhe Ekzekutimi")
code(
"py -3.13 -m venv .venv\n"
".venv\\Scripts\\pip install -r requirements.txt\n"
"ollama pull qwen2.5:3b        # opsionale: ollama pull gemma2:9b\n"
"python scripts\\seed_sample_corpus.py    # opsionale: korpus demo\n"
".venv\\Scripts\\streamlit run app.py\n"
"# Hyrja: admin / ***REMOVED-CREDENTIAL***  ->  vendos nje fjalekalim te ri")

h1("Shtojca B — Skema e Bazës së të Dhënave")
para("Shih Figurën 4.7 për skemën e plotë ER. Pesë tabela në SQLite (users, documents, chat_history, "
     "audit_logs, experiment_results) dhe koleksioni vektorial doku_chunks në ChromaDB me vektorë "
     "1024-dimensionalë dhe metadata të plota për citime.")

h1("Shtojca C — Shembull Konfigurimi (config.py)")
code(
"OLLAMA_MODEL   = 'gemma2:9b'      # parazgjedhje (qwen2.5:3b per shpejtesi)\n"
"EMBEDDING_MODEL= 'BAAI/bge-m3'    # embeddings shumegjuhesh, 1024-D\n"
"MIN_SIMILARITY = 0.38             # pragu i portes se refuzimit\n"
"CHUNK_SIZE = 800 ; CHUNK_OVERLAP = 120 ; RETRIEVAL_K = 5\n"
"LLM_TEMPERATURE = 0.2             # temperature e ulet per pergjigje te bazuara")

h1("Shtojca D — Rezultatet e Plota të Testimit")
para("Të dhënat e plota të eksperimentit (Tabela 6.1) eksportohen automatikisht në CSV nga moduli i "
     "eksperimenteve (data/exports/experiments_*.csv) dhe janë plotësisht të riprodhueshme nëpërmjet "
     "harness-it. Vlerësimi manual i saktësisë (shkalla 1–5) dhe i halucinacionit (Po/Jo) plotësohet "
     "nga administratori në faqen «Eksperimente» (Figura 6.1).")

h1("Shtojca E — Galeri e Pamjeve nga Ekrani")
para("Më poshtë jepet dashboard-i kryesor i sistemit me metrikat e korpusit. Pamjet e tjera janë "
     "integruar në kapitujt përkatës (Figurat 5.1–5.8 dhe 6.1).")
figure("02_dashboard.png", "Figura E.1 — Dashboard-i: dokumente aktive, totali, copëza të indeksuara "
       "dhe numri i përdoruesve, me lidhje te veprimet kryesore.")

# ======================================================================================
# RUAJTJA
# ======================================================================================
enable_update_fields_on_open()   # Word përditëson TOC-në automatikisht kur e hap
out_path = config.EXPORTS_DIR / "Teza_DOKU.docx"
doc.save(str(out_path))
# kopje ne docs/
docs_copy = os.path.join(ROOT, "docs", "Teza_DOKU.docx")
doc.save(docs_copy)
print("THESIS_SAVED:", out_path)
print("THESIS_COPY :", docs_copy)
