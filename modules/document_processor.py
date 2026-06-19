"""Document processing: extract text from PDF (PyMuPDF, per page) or Word/.docx
(python-docx), validate that the file actually has selectable text (no OCR in the
base version), and split into overlapping chunks carrying the source page number."""
from dataclasses import dataclass

import fitz  # PyMuPDF
from docx import Document as DocxDocument

import config


@dataclass
class Chunk:
    text: str
    page_number: int   # 1-based
    chunk_index: int


class NoExtractableTextError(Exception):
    """Raised when a PDF has (almost) no selectable text — likely a scan."""


def extract_pages(pdf_path) -> list[str]:
    pages = []
    with fitz.open(pdf_path) as doc:
        for page in doc:
            pages.append(page.get_text("text"))
    return pages


def extract_docx(docx_path) -> list[str]:
    """Word has no fixed pages; return the whole text as a single 'page'."""
    doc = DocxDocument(docx_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(c.text for c in row.cells if c.text.strip()))
    return ["\n".join(parts)]


def _pages_for(path) -> list[str]:
    name = str(path).lower()
    if name.endswith(".pdf"):
        return extract_pages(path)
    if name.endswith(".docx"):
        return extract_docx(path)
    raise ValueError("Format i pambështetur. Lejohen vetëm PDF ose DOCX.")


def extract_text(path) -> str:
    """Full plain text of a document (PDF or DOCX) — used for summaries."""
    return "\n".join(_pages_for(path))


def render_pdf_images(pdf_path, max_pages: int = 10, zoom: float = 1.5) -> list[bytes]:
    """Render PDF pages to PNG bytes for an in-app preview (no extra deps)."""
    images = []
    with fitz.open(pdf_path) as doc:
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            images.append(pix.tobytes("png"))
    return images


def validate_has_text(pages: list[str], min_chars: int = 100) -> None:
    if sum(len(p.strip()) for p in pages) < min_chars:
        raise NoExtractableTextError(
            "Ky PDF nuk përmban tekst të lexueshëm (ndoshta është i skanuar). "
            "Versioni bazë nuk mbështet OCR."
        )


def chunk_text(text: str, size: int, overlap: int) -> list[str]:
    """Character-window chunking with overlap, breaking near whitespace."""
    text = text.strip()
    if not text:
        return []
    chunks, start, n = [], 0, len(text)
    while start < n:
        end = min(start + size, n)
        if end < n:
            ws = text.rfind(" ", start + overlap, end)
            if ws != -1 and ws > start:
                end = ws
        piece = text[start:end].strip()
        if piece:
            chunks.append(piece)
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return chunks


def process_document(path) -> tuple[list[Chunk], int]:
    """Extract, validate, and chunk a PDF or DOCX. Returns (chunks, n_pages).
    Raises NoExtractableTextError if the document has no usable text."""
    pages = _pages_for(path)
    validate_has_text(pages)
    chunks, idx = [], 0
    for page_no, page_text in enumerate(pages, start=1):
        for piece in chunk_text(page_text, config.CHUNK_SIZE, config.CHUNK_OVERLAP):
            chunks.append(Chunk(text=piece, page_number=page_no, chunk_index=idx))
            idx += 1
    if not chunks:
        raise NoExtractableTextError("Nuk u nxor asnjë tekst nga dokumenti.")
    return chunks, len(pages)
