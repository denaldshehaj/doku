"""Word (.docx) export for answers and summaries. Files are saved to
data/exports/ with safe, unique filenames and returned for the Streamlit download
button."""
from datetime import datetime

from docx import Document

import config

VERIFY_NOTE = ("Kjo përgjigje është gjeneruar automatikisht nga sistemi DOKU duke "
               "përdorur RAG + LLM lokal dhe duhet verifikuar me dokumentet origjinale.")
SUMMARY_VERIFY_NOTE = ("Kjo përmbledhje është gjeneruar automatikisht nga sistemi DOKU "
                       "dhe duhet verifikuar me dokumentin origjinal.")


def _stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def export_answer_to_docx(user_id, username, question, answer, sources,
                          response_time=None) -> str:
    doc = Document()
    doc.add_heading("Përgjigje e Gjeneruar nga Sistemi DOKU", level=1)
    doc.add_paragraph(f"Data dhe ora: {datetime.now():%Y-%m-%d %H:%M}")
    doc.add_paragraph(f"Përdoruesi: {username}")
    if response_time is not None:
        doc.add_paragraph(f"Koha e përgjigjes: {response_time}s")

    doc.add_heading("Pyetja", level=2)
    doc.add_paragraph(question)

    doc.add_heading("Përgjigjja", level=2)
    for para in (answer or "").split("\n"):
        doc.add_paragraph(para)

    if sources:
        doc.add_heading("Burimet", level=2)
        for s in sources:
            doc.add_paragraph(
                f"[{s['n']}] {s['filename']}, {s.get('document_type','')}, "
                f"{s.get('institution','')}, faqe {s['page']} — {s.get('fragment','')}",
                style="List Number")

    doc.add_heading("Shënim", level=2)
    doc.add_paragraph(VERIFY_NOTE)

    out = config.EXPORTS_DIR / f"answer_{_stamp()}_{user_id}.docx"
    doc.save(str(out))
    return str(out)


def export_summary_to_docx(doc_meta, summary, fmt="") -> str:
    doc = Document()
    doc.add_heading("Përmbledhje Dokumenti e Gjeneruar nga Sistemi DOKU", level=1)
    doc.add_paragraph(f"Data e gjenerimit: {datetime.now():%Y-%m-%d %H:%M}")

    doc.add_heading("Të dhënat e dokumentit", level=2)
    doc.add_paragraph(f"Titulli: {doc_meta.get('title','')}")
    doc.add_paragraph(f"Skedari PDF: {doc_meta.get('filename','')}")
    doc.add_paragraph(f"Institucioni burimor: {doc_meta.get('institution','')}")
    doc.add_paragraph(f"Tipi i dokumentit: {doc_meta.get('document_type','')}")
    doc.add_paragraph(f"Viti: {doc_meta.get('year','')}")
    if fmt:
        doc.add_paragraph(f"Formati i përmbledhjes: {fmt}")

    doc.add_heading("Përmbledhja", level=2)
    for para in (summary or "").split("\n"):
        doc.add_paragraph(para)

    doc.add_heading("Shënim", level=2)
    doc.add_paragraph(SUMMARY_VERIFY_NOTE)

    out = config.EXPORTS_DIR / f"summary_{_stamp()}_{doc_meta.get('id','x')}.docx"
    doc.save(str(out))
    return str(out)
