"""Word (.docx) export for answers and summaries, including citations."""
from datetime import datetime
from io import BytesIO

from docx import Document


def build_docx(title: str, body: str, citations: list[dict] | None = None) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"Gjeneruar nga DOKU — {datetime.now():%Y-%m-%d %H:%M}").italic = True

    for para in body.split("\n"):
        doc.add_paragraph(para)

    if citations:
        doc.add_heading("Burimet", level=2)
        for c in citations:
            doc.add_paragraph(
                f"[{c['n']}] {c['title']} — faqe {c['page']} "
                f"(ngjashmëria: {c['score']})",
                style="List Bullet",
            )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
